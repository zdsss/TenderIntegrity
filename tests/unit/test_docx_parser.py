"""DocxParser 单元测试（使用 python-docx 构建测试文档）"""
import pytest
from pathlib import Path
import tempfile

from src.document.docx_parser import DocxParser, HEADING_MARKER


def _create_test_docx(tmp_path: Path) -> Path:
    """创建含段落、表格、Heading 样式的测试 .docx 文件"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        pytest.skip("python-docx 未安装，跳过 DocxParser 测试")

    doc = Document()

    # Heading 段落
    doc.add_heading("第一章 技术方案", level=1)

    # 普通段落
    doc.add_paragraph("本方案采用先进的医疗影像技术，配置高分辨率探头。")

    # 表格（带表头）
    table = doc.add_table(rows=3, cols=3)
    headers = ["姓名", "岗位", "年限"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "张三"
    table.rows[1].cells[1].text = "项目经理"
    table.rows[1].cells[2].text = "10年"
    table.rows[2].cells[0].text = "李四"
    table.rows[2].cells[1].text = "技术总监"
    table.rows[2].cells[2].text = "8年"

    # Heading 段落（第二章）
    doc.add_heading("第二章 商务条款", level=1)

    # 普通段落
    doc.add_paragraph("报价按合同约定执行，含税单价，须加盖公章。")

    fpath = tmp_path / "test_doc.docx"
    doc.save(str(fpath))
    return fpath


def test_docx_parser_returns_full_text(tmp_path):
    fpath = _create_test_docx(tmp_path)
    parser = DocxParser()
    text, page_map = parser.parse(fpath)

    assert isinstance(text, str)
    assert len(text) > 0
    assert page_map == {}  # Word 不提供精确页码


def test_docx_parser_heading_marker(tmp_path):
    """Heading 样式段落应带 §H§ 前缀"""
    fpath = _create_test_docx(tmp_path)
    parser = DocxParser()
    text, _ = parser.parse(fpath)

    lines = [p.strip() for p in text.split("\n\n") if p.strip()]
    heading_lines = [l for l in lines if l.startswith(HEADING_MARKER)]
    assert len(heading_lines) >= 1, "应有至少一行带 §H§ 标记的 Heading"
    assert any("技术方案" in l or "商务条款" in l for l in heading_lines)


def test_docx_parser_table_structured(tmp_path):
    """表格应以 [表格] 前缀输出，含 key:value 格式"""
    fpath = _create_test_docx(tmp_path)
    parser = DocxParser()
    text, _ = parser.parse(fpath)

    lines = [p.strip() for p in text.split("\n\n") if p.strip()]
    table_lines = [l for l in lines if l.startswith("[表格]")]
    assert len(table_lines) >= 2, "应有至少 2 行表格数据"
    # key:value 格式
    assert any("姓名:张三" in l or "姓名:李四" in l for l in table_lines)


def test_docx_parser_table_header_marked(tmp_path):
    """表格标题行应带 §H§[表格标题] 前缀"""
    fpath = _create_test_docx(tmp_path)
    parser = DocxParser()
    text, _ = parser.parse(fpath)

    lines = [p.strip() for p in text.split("\n\n") if p.strip()]
    header_lines = [
        l for l in lines
        if l.startswith(HEADING_MARKER) and "[表格标题]" in l
    ]
    assert len(header_lines) >= 1, "应有至少一个表格标题行"
    assert any("姓名" in l and "岗位" in l for l in header_lines)


def test_docx_parser_document_order(tmp_path):
    """段落和表格应按文档原始顺序排列（Heading1 在表格之前）"""
    fpath = _create_test_docx(tmp_path)
    parser = DocxParser()
    text, _ = parser.parse(fpath)

    lines = [p.strip() for p in text.split("\n\n") if p.strip()]
    # 找到第一章 Heading 和表格的位置
    heading_idx = next(
        (i for i, l in enumerate(lines) if HEADING_MARKER in l and "技术方案" in l), -1
    )
    table_idx = next(
        (i for i, l in enumerate(lines) if "[表格]" in l), -1
    )
    assert heading_idx >= 0, "应找到第一章标题"
    assert table_idx >= 0, "应找到表格行"
    # 第一章在表格之前
    assert heading_idx < table_idx, "第一章标题应在表格之前（文档顺序）"


def test_docx_parser_file_not_found():
    parser = DocxParser()
    with pytest.raises(FileNotFoundError):
        parser.parse("/nonexistent/path/file.docx")


def test_docx_parser_empty_paragraphs(tmp_path):
    """空段落不应出现在输出中"""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx 未安装")

    doc = Document()
    doc.add_paragraph("")  # 空段落
    doc.add_paragraph("有效内容段落，包含足够的文字。")
    doc.add_paragraph("")  # 空段落

    fpath = tmp_path / "empty_paras.docx"
    doc.save(str(fpath))

    parser = DocxParser()
    text, _ = parser.parse(fpath)

    lines = [p.strip() for p in text.split("\n\n") if p.strip()]
    # 不应有空行
    for line in lines:
        assert line.strip() != ""
