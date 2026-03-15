"""Word 文档解析（python-docx）— 按 XML body 原始顺序，结构化表格"""
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# 内部标记：表示该行为 Heading 样式段落（由 ChunkSplitter 识别后剥离）
HEADING_MARKER = "§H§"


class DocxParser:
    """
    使用 python-docx 解析 .docx 文件。

    改进：
    - 按 XML body 顺序遍历，保留段落与表格的文档原始顺序
    - 检测 Word Heading 样式，用 §H§ 前缀标记
    - 表格结构化处理：首行为表头时生成 key:value 行
    """

    def parse(self, file_path: str | Path) -> tuple[str, dict[int, int]]:
        """
        解析 .docx 文件

        Returns:
            (full_text, page_map)
            page_map: Word 无精确页码，返回空 dict
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: uv add python-docx")

        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = Document(str(file_path))

        # 构建 XML element → docx 对象的映射（保留 part 上下文）
        para_map: dict = {}
        for para in doc.paragraphs:
            para_map[id(para._element)] = para

        table_map: dict = {}
        for table in doc.tables:
            table_map[id(table._element)] = table

        parts: list[str] = []
        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para = para_map.get(id(child))
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue
                style_name = self._get_style_name(para)
                if (
                    "Heading" in style_name
                    or "heading" in style_name
                    or "标题" in style_name
                ):
                    parts.append(f"{HEADING_MARKER}{text}")
                else:
                    parts.append(text)

            elif tag == "tbl":
                table = table_map.get(id(child))
                if table is None:
                    continue
                table_parts = self._process_table(table)
                parts.extend(table_parts)

        full_text = "\n\n".join(parts)
        logger.info(
            f"DOCX 解析完成: {file_path.name}, {len(parts)} 段落/表格行, "
            f"{len(full_text)} 字符"
        )
        return full_text, {}

    # ── 辅助方法 ──────────────────────────────────────────────────

    @staticmethod
    def _get_style_name(para) -> str:
        try:
            if para.style:
                return para.style.name or ""
        except Exception:
            pass
        return ""

    # ── 表格结构化处理 ──────────────────────────────────────────────

    def _process_table(self, table) -> list[str]:
        """将表格转换为结构化文本行"""
        result: list[str] = []
        rows = list(table.rows)
        if not rows:
            return result

        # 去重（合并单元格在 python-docx 中会返回重复对象）
        def dedup_cells(row) -> list[str]:
            seen_ids: set[int] = set()
            cells: list[str] = []
            for cell in row.cells:
                cid = id(cell._tc)
                if cid not in seen_ids:
                    seen_ids.add(cid)
                    cells.append(cell.text.strip())
            return cells

        first_row_cells = dedup_cells(rows[0])
        # 判断首行是否可作表头：全部单元格非空且 >1 列
        has_header = len(first_row_cells) > 1 and all(c for c in first_row_cells)

        if has_header:
            headers = first_row_cells
            # 表头行作为一个 chunk（标记为 HEADING，保证章节归属）
            result.append(f"{HEADING_MARKER}[表格标题] {' | '.join(headers)}")
            for row in rows[1:]:
                cells = dedup_cells(row)
                if not any(cells):
                    continue
                if len(headers) == len(cells):
                    row_text = " | ".join(
                        f"{h}:{v}" for h, v in zip(headers, cells) if v
                    )
                else:
                    row_text = " | ".join(c for c in cells if c)
                if row_text:
                    result.append(f"[表格] {row_text}")
        else:
            for row in rows:
                cells = dedup_cells(row)
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    result.append(f"[表格] {row_text}")

        return result
